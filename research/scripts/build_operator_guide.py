#!/usr/bin/env python3
"""Build the True North Map research skills operator guide from Markdown."""

from pathlib import Path
import re
import subprocess
import sys
import tempfile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image


NAVY = RGBColor(18, 42, 62)
BLUE = RGBColor(29, 92, 122)
GOLD = RGBColor(197, 145, 53)
SLATE = RGBColor(74, 88, 99)
LIGHT_BLUE = "EAF2F6"


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def set_paragraph_border(paragraph, color: str = "C59135") -> None:
    properties = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    properties.append(borders)


def add_inline_markdown(paragraph, text: str) -> None:
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Menlo"
            run.font.size = Pt(10.5)
            run.font.color.rgb = BLUE
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(12)
    normal.font.color.rgb = NAVY
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color in [
        ("Title", 25, NAVY),
        ("Heading 1", 19, NAVY),
        ("Heading 2", 14, BLUE),
        ("Heading 3", 12, SLATE),
    ]:
        style = document.styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(10 if name != "Title" else 0)
        style.paragraph_format.space_after = Pt(5)

    if "Code Block" not in [style.name for style in document.styles]:
        code = document.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = document.styles["Code Block"]
    code.font.name = "Menlo"
    code.font.size = Pt(9.5)
    code.font.color.rgb = NAVY
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.right_indent = Inches(0.18)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(6)

    header = section.header.paragraphs[0]
    header.text = "TRUE NORTH MAP  /  RESEARCH OPERATIONS"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header.runs[0]
    header_run.font.name = "Aptos"
    header_run.font.size = Pt(8)
    header_run.font.bold = True
    header_run.font.color.rgb = GOLD

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("truenorthmap.ca   |   Private review-first research pipeline   |   ")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)
    for run in footer.runs:
        run.font.name = "Aptos"
        run.font.size = Pt(8)
        run.font.color.rgb = SLATE


def render_mermaid(source: str, output_path: Path, work_dir: Path) -> None:
    source_path = work_dir / f"{output_path.stem}.mmd"
    config_path = work_dir / "puppeteer-config.json"
    source_path.write_text(source, encoding="utf-8")
    if not config_path.exists():
        config_path.write_text(
            '{"executablePath":"/Applications/Google Chrome.app/Contents/MacOS/Google Chrome","args":["--no-sandbox"]}',
            encoding="utf-8",
        )
    command = [
        "npx",
        "-y",
        "@mermaid-js/mermaid-cli",
        "-p",
        str(config_path),
        "-i",
        str(source_path),
        "-o",
        str(output_path),
        "-b",
        "transparent",
        "-s",
        "2",
    ]
    result = subprocess.run(command, capture_output=True, text=True, check=False)
    if result.returncode != 0 or not output_path.exists():
        detail = (result.stderr or result.stdout).strip()
        raise RuntimeError(f"Mermaid rendering failed: {detail}")


def add_mermaid_figure(document: Document, source: str, output_path: Path, work_dir: Path, figure_number: int) -> None:
    render_mermaid(source, output_path, work_dir)
    with Image.open(output_path) as rendered:
        aspect = rendered.width / rendered.height
    width_inches = 6.75
    height_inches = width_inches / aspect
    if height_inches > 7.65:
        height_inches = 7.65
        width_inches = height_inches * aspect
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.add_run().add_picture(str(output_path), width=Inches(width_inches), height=Inches(height_inches))
    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(8)
    run = caption.add_run(f"Workflow {figure_number}")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = SLATE


def build(markdown_path: Path, docx_path: Path) -> None:
    document = Document()
    configure_document(document)

    in_code = False
    code_lines: list[str] = []
    code_language = ""
    first_h1 = True
    figure_number = 0

    with tempfile.TemporaryDirectory(prefix="tnm-operator-guide-", dir=docx_path.parent) as temp_name:
        mermaid_dir = Path(temp_name)
        for raw_line in markdown_path.read_text(encoding="utf-8").splitlines():
            line = raw_line.rstrip()
            if line.startswith("```"):
                if in_code:
                    if code_language == "mermaid":
                        figure_number += 1
                        output_path = mermaid_dir / f"workflow-{figure_number:02d}.png"
                        add_mermaid_figure(document, "\n".join(code_lines), output_path, mermaid_dir, figure_number)
                    else:
                        paragraph = document.add_paragraph(style="Code Block")
                        paragraph.add_run("\n".join(code_lines))
                        set_cell = OxmlElement("w:shd")
                        set_cell.set(qn("w:fill"), LIGHT_BLUE)
                        paragraph._p.get_or_add_pPr().append(set_cell)
                    code_lines = []
                    code_language = ""
                else:
                    code_language = line[3:].strip().lower()
                in_code = not in_code
                continue
            if in_code:
                code_lines.append(line)
                continue
            if not line:
                continue

            if line.startswith("# "):
                title = line[2:]
                if first_h1:
                    paragraph = document.add_paragraph(style="Title")
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = paragraph.add_run(title)
                    run.font.color.rgb = NAVY
                    set_paragraph_border(paragraph)
                    subtitle = document.add_paragraph()
                    subtitle.paragraph_format.space_after = Pt(12)
                    subtitle_run = subtitle.add_run("Multi-source discovery and refresh -> enriched private candidates -> human review and publication")
                    subtitle_run.italic = True
                    subtitle_run.font.color.rgb = SLATE
                    first_h1 = False
                else:
                    document.add_page_break()
                    paragraph = document.add_paragraph(title, style="Heading 1")
                    set_paragraph_border(paragraph)
                continue
            if line.startswith("## "):
                document.add_paragraph(line[3:], style="Heading 2")
                continue
            if line.startswith("### "):
                document.add_paragraph(line[4:], style="Heading 3")
                continue
            if line.startswith("- "):
                paragraph = document.add_paragraph(style="List Bullet")
                paragraph.paragraph_format.space_after = Pt(3)
                add_inline_markdown(paragraph, line[2:])
                continue
            numbered = re.match(r"^(\d+)\.\s+(.*)$", line)
            if numbered:
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.left_indent = Inches(0.22)
                paragraph.paragraph_format.first_line_indent = Inches(-0.22)
                paragraph.paragraph_format.space_after = Pt(3)
                number_run = paragraph.add_run(f"{numbered.group(1)}. ")
                number_run.bold = True
                add_inline_markdown(paragraph, numbered.group(2))
                continue
            if line.startswith("> "):
                table = document.add_table(rows=1, cols=1)
                table.autofit = True
                cell = table.cell(0, 0)
                set_cell_shading(cell, LIGHT_BLUE)
                paragraph = cell.paragraphs[0]
                paragraph.paragraph_format.space_after = Pt(0)
                run = paragraph.add_run(line[2:])
                run.italic = True
                run.font.color.rgb = BLUE
                continue

            paragraph = document.add_paragraph()
            add_inline_markdown(paragraph, line)

    core_properties = document.core_properties
    core_properties.title = "True North Map Research Pipeline Skills - Operator Guide"
    core_properties.subject = "Review-first autonomous research pipeline operations"
    core_properties.author = "True North Map"
    core_properties.keywords = "True North Map, research, ingestion, skills, Admin Review"
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(docx_path)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        raise SystemExit("Usage: build_operator_guide.py <source.md> <output.docx>")
    build(Path(sys.argv[1]), Path(sys.argv[2]))
